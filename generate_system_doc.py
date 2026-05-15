"""
Generate a Microsoft Word document summarising the GeoAI Real-Estate Intelligence Platform.
Run: python generate_system_doc.py
Output: GeoAI_RealEstate_System_Overview.docx  (project root)
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_col_width(table, col_idx: int, width_cm: float):
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)

def add_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    if color:
        for run in h.runs:
            run.font.color.rgb = RGBColor(*color)
    return h

def add_bullet(doc, text, bold_prefix: str = None, indent=1):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(indent * 0.75)
    if bold_prefix:
        run = p.add_run(bold_prefix + " ")
        run.bold = True
    p.add_run(text)
    return p

def add_kv_table(doc, rows: list[tuple[str, str]], hdr_color="1F4E79", alt_color="D6E4F0"):
    table = doc.add_table(rows=len(rows) + 1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Property"
    hdr_cells[1].text = "Detail"
    for cell in hdr_cells:
        set_cell_bg(cell, hdr_color)
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
    for i, (k, v) in enumerate(rows):
        row = table.rows[i + 1]
        row.cells[0].text = k
        row.cells[1].text = v
        bg = alt_color if i % 2 == 0 else "FFFFFF"
        set_cell_bg(row.cells[0], bg)
        set_cell_bg(row.cells[1], bg)
    set_col_width(table, 0, 6)
    set_col_width(table, 1, 11)
    doc.add_paragraph()

def add_metric_table(doc, headers: list[str], data: list[list[str]], hdr_color="1F4E79", alt_color="D6E4F0"):
    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.style = "Table Grid"
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        set_cell_bg(cell, hdr_color)
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
    for i, row_data in enumerate(data):
        row = table.rows[i + 1]
        for j, val in enumerate(row_data):
            row.cells[j].text = val
            bg = alt_color if i % 2 == 0 else "FFFFFF"
            set_cell_bg(row.cells[j], bg)
    doc.add_paragraph()


# ─────────────────────────────────────────────────────────────────────────────
# Build document
# ─────────────────────────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# Default paragraph style
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# ── Cover ────────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("GeoAI Real-Estate Intelligence Platform")
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = RGBColor(31, 78, 121)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("System Overview & Technical Reference")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(70, 130, 180)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run(f"Document prepared: {datetime.date.today().strftime('%B %d, %Y')}").italic = True

doc.add_page_break()

# ── 1. Executive Summary ─────────────────────────────────────────────────────
add_heading(doc, "1. Executive Summary", level=1, color=(31, 78, 121))

doc.add_paragraph(
    "The GeoAI Real-Estate Intelligence Platform is a full-stack AI/ML system built to "
    "automate property valuation, forecast price trajectories, and generate actionable "
    "investment intelligence for the Indian real estate market — initially covering the "
    "National Capital Region (NCR) and Jaipur. The platform ingests raw listing data from "
    "multiple scrapers, applies automated feature engineering, trains production-grade "
    "machine learning models, and exposes all results through a REST API consumed by a "
    "modern React dashboard."
)

doc.add_paragraph(
    "The system covers three distinct property segments — Apartments, Builder Floors, and "
    "Plots/Land — with separate ML pipelines optimised for each segment's unique valuation "
    "drivers. Beyond simple price estimation, the platform delivers forecast intelligence "
    "(5-year quarterly horizon), locality-level market heat indices, ROI calculators, "
    "buy/watch/avoid recommendations, and explainable-AI narratives for every prediction."
)

doc.add_paragraph()

# ── 2. High-Level Architecture ───────────────────────────────────────────────
add_heading(doc, "2. High-Level Architecture", level=1, color=(31, 78, 121))

doc.add_paragraph(
    "The platform is organised into four layers that together form an end-to-end ML system:"
)

add_bullet(doc, "Raw CSV data from two listing scrapers (HO and MB) stored in real_estate_data/", "Data Layer —")
add_bullet(doc, "Python pipeline: ingestion → transformation → model training, tracked with MLflow", "ML Pipeline —")
add_bullet(doc, "FastAPI backend (api/main.py) serving predictions, forecasts, and intelligence endpoints", "API Layer —")
add_bullet(doc, "React + Vite single-page application consuming all API endpoints in real time", "Frontend Layer —")

doc.add_paragraph()

add_heading(doc, "2.1 Data Flow", level=2)

flow_steps = [
    ("Step 1", "Raw listing data (ho_raw_data.csv, mb_raw_data.csv) scraped from online portals"),
    ("Step 2", "DataIngestion component merges both sources into a single merged.csv (~300 K rows)"),
    ("Step 3", "DataTransformation applies NLP parsing of descriptions, circle-rate joining, locality matching, and full feature engineering → cleaned.csv"),
    ("Step 4", "Segment-specific trainers (AptModelTrainer, BfModelTrainer, PlotModelTrainer) read cleaned.csv and produce versioned model bundles"),
    ("Step 5", "Forecast training (app_apt_train.py / app_bf_train.py / app_plot_train.py) runs the rho-decay engine and writes opt/{segment}/ artifacts"),
    ("Step 6", "FastAPI loads all models at startup in parallel (ThreadPoolExecutor) and serves live predictions"),
    ("Step 7", "React frontend fetches /meta/options on boot, then calls prediction and intelligence endpoints on demand"),
]

add_metric_table(doc,
    ["Step", "Description"],
    flow_steps,
    hdr_color="1F4E79",
    alt_color="D6E4F0"
)

# ── 3. ML Pipelines ──────────────────────────────────────────────────────────
add_heading(doc, "3. Machine Learning Pipelines", level=1, color=(31, 78, 121))

doc.add_paragraph(
    "Three independent ML pipelines exist — one per property segment. "
    "Each pipeline follows the same Pattern: data filtering → outlier removal → "
    "feature engineering → cross-validated model selection → evaluation → artifact persistence."
)

# 3.1 Apartment & Builder Floor
add_heading(doc, "3.1 Apartment & Builder Floor Models", level=2)

doc.add_paragraph(
    "Both models use the same mathematical target and feature construction strategy; "
    "they differ only in the features specific to multi-storey apartments (floor level, "
    "property segment tier) versus builder floors (ground-level floor attributes)."
)

add_heading(doc, "Target Variable Design", level=3)
doc.add_paragraph(
    "Instead of predicting raw price-per-sqft directly, the models predict the ratio of "
    "price-per-sqft to the government-mandated circle rate, then apply a log1p transform "
    "to compress the right-skewed distribution:"
)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("target  =  log₁ₚ ( price_per_sqft / circle_rate )")
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(31, 78, 121)

doc.add_paragraph(
    "At inference the prediction is inverted: "
    "predicted_ppsf = expm1(model_output) × circle_rate. "
    "This design makes the model invariant to city-level price level differences "
    "(Noida vs Gurgaon vs Jaipur) while still capturing micro-locality premiums."
)

add_heading(doc, "Key Features", level=3)
add_kv_table(doc, [
    ("Structural",      "BHK, bathrooms, balconies, covered area (sqft)"),
    ("Amenities",       "is_parking, is_pool, is_main_road, is_garden_park, is_gated, is_corner"),
    ("Property Age",    "OHE buckets: New Construction / <5 yrs / 5–10 yrs / 10–20 yrs / >20 yrs"),
    ("Furnishing",      "OHE: Furnished / Semi-Furnished / Unfurnished"),
    ("Facing",          "OHE: North / South / East / West / North-East etc."),
    ("Floor (APT)",     "floor_low / floor_medium / floor_high dummies; is_ground_floor; is_top_floor"),
    ("Floor (BF)",      "current_floor, total_floors, is_ground_floor, is_top_floor, is_basement"),
    ("Spatial",         "Voronoi cell dummy (KMeans on lat/lon) + dist_to_seed (km)"),
    ("Circle Rate",     "Government circle rate for the locality (INR/sqft) — key locality proxy"),
    ("Locality Tier",   "locality_tier_ord (ordinal 0–3) + OHE dummies (budget/mid/high/premium/luxury)"),
], hdr_color="1F4E79")

add_heading(doc, "Training Process", level=3)
add_bullet(doc, "IsolationForest outlier removal after train/test split (contamination = 5%)")
add_bullet(doc, "5-fold cross-validation to select the best RandomForestRegressor hyperparameters")
add_bullet(doc, "Final model fitted on full training set")
add_bullet(doc, "Evaluation in real units: MAE, MAPE, R² computed on back-calculated price_per_sqft")
add_bullet(doc, "Model bundle saved as joblib: {model, kmeans, locality_tier_map, features, target}")

doc.add_paragraph()

# 3.2 Plot Model
add_heading(doc, "3.2 Plot / Land Model", level=2)

doc.add_paragraph(
    "The plot model uses a completely separate data pipeline (PlotDataIngestion, "
    "PlotDataTransformation) because plot listings have fundamentally different attributes "
    "— there is no BHK, no floor, no furnishing. The target is log1p(price_per_sqft) "
    "directly (circle rate is retained as a feature, not as a normaliser)."
)

add_heading(doc, "Key Features", level=3)
add_kv_table(doc, [
    ("Area",              "log_plot_area = log1p(area_sqft) — reduces skew"),
    ("Circle Rate",       "Government circle rate (INR/sqft)"),
    ("Spatial",           "KMeans cluster dummies (c_0 … c_N) + dist_to_center (km)"),
    ("Road Network",      "closest_distance_NH_km, closest_distance_SH_km, closest_distance_MDR_km (highway proximity)"),
    ("Road Width",        "road_width_upto_9m / road_width_9_to_18m / road_width_18_plus (OHE)"),
    ("Plot Shape/Type",   "is_rectangular, is_corner, is_park_facing, is_gated, has_boundary_wall"),
    ("Usage Type",        "OHE: Residential / Commercial / Industrial / Mixed"),
    ("Facing",            "OHE: North / South / East / West"),
    ("Coordinates",       "latitude, longitude (direct — tree models handle non-linearity)"),
], hdr_color="244061")

doc.add_paragraph()

# ── 4. Forecast Intelligence Engine ─────────────────────────────────────────
add_heading(doc, "4. Forecast Intelligence Engine", level=1, color=(31, 78, 121))

doc.add_paragraph(
    "The forecast engine produces per-property 5-year quarterly price trajectories using a "
    "custom rho-decay blending model. It reconciles two signals — the ML model's intrinsic "
    "price estimate (P_i) and the listing price (L_i) — and anchors both to a locality-level "
    "growth index (I_l,t) derived from macro indicators (GSDP, RBI repo rate, home loan rates)."
)

add_heading(doc, "4.1 Core Formula", level=2)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("F_{i,t}  =  [ ρ_t · L_i  +  (1 − ρ_t) · P_i ]  ×  I_{l,t}")
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(31, 78, 121)

doc.add_paragraph()
add_kv_table(doc, [
    ("F_{i,t}",  "Forecasted price-per-sqft for property i at quarter t"),
    ("L_i",      "Listing price-per-sqft (observed market price)"),
    ("P_i",      "ML model intrinsic value estimate"),
    ("I_{l,t}",  "Locality growth index at quarter t (derived from macro data)"),
    ("ρ_t",      "Time-varying blend weight — starts at ρ₀, decays geometrically"),
    ("ρ₀",       "clip(0.25 + 0.30·C_i + 0.20·U_i − 0.25·M_i, 0.10, 0.70)"),
    ("ρ_t decay","ρ_t = ρ₀ × 0.90^t  (10% decay per quarter; converges to model at t→∞)"),
    ("C_i",      "Comparable-support score — how many similar listings anchor this price"),
    ("U_i",      "Uniqueness score — how differentiated the property is"),
    ("M_i",      "Model confidence — how well the ML model fits this price tier"),
], hdr_color="1F4E79")

doc.add_paragraph(
    "Intuition: in the short term the listing price dominates (ρ is high); "
    "over time the model's intrinsic value takes over as temporary market sentiment fades. "
    "The locality index I_l,t embeds macroeconomic growth so the absolute price level "
    "evolves with the broader economy."
)

add_heading(doc, "4.2 Forecast Outputs per Property", level=2)
add_bullet(doc, "Quarter-by-quarter forecast table (date, P_i, L_i, ρ₀, ρ_t, I_lt, forecast_ppsf)")
add_bullet(doc, "Year-over-year (YoY) % appreciation series for the property")
add_bullet(doc, "YoY % appreciation for the locality (aggregate of all properties)")
add_bullet(doc, "Price distribution band (p25 / median / p75 / min / max) across all locality properties")
add_bullet(doc, "Historical price trend (from training data) spliced with forecast series")
add_bullet(doc, "Math validation: linear vs. log-space formula cross-check with mean/max absolute error")

doc.add_paragraph()

# ── 5. Intelligence Tabs ─────────────────────────────────────────────────────
add_heading(doc, "5. Intelligence Modules", level=1, color=(31, 78, 121))

add_heading(doc, "5.1 Locality Intelligence", level=2)
doc.add_paragraph(
    "For any selected locality the platform computes a comprehensive profile from listing data:"
)
add_bullet(doc, "Median price-per-sqft and its trend over time")
add_bullet(doc, "Listing count (demand/supply proxy)")
add_bullet(doc, "Affordability score = clip( circle_rate / median_ppsf × 50, 0, 100 )")
add_bullet(doc, "Forecast appreciation (YoY %) from today using the rho-decay engine")
add_bullet(doc, "Volatility: standard deviation of QoQ % changes of forecast prices")
add_bullet(doc, "Combined historical + forecast price trend chart")
add_bullet(doc, "BHK distribution and price-per-BHK analysis")

doc.add_paragraph()

add_heading(doc, "5.2 Market Intelligence", level=2)
doc.add_paragraph(
    "The market intelligence module ingests all raw listing events (new / updated / removed) "
    "and computes supply/demand metrics at the locality × scrape-date granularity:"
)

add_metric_table(doc,
    ["Metric", "Formula / Definition"],
    [
        ("Active Supply Stock",    "Count of unique active listing IDs at each scrape date"),
        ("Absorption Rate",        "Listings vanished to next scrape ÷ active supply (demand proxy)"),
        ("Stale Inventory Share",  "Listings with days_on_market > 90 days ÷ active supply"),
        ("New Supply Velocity",    "New listings at scrape date ÷ active supply"),
        ("Price Cut Frequency",    "Updated listings with price reduction ÷ total updated listings"),
        ("Price Hike Frequency",   "Updated listings with price increase ÷ total updated listings"),
        ("Median Days on Market",  "Median of (scrape_date − posting_date) across all active listings"),
        ("Supply Pressure Index",  "0.40·S_active + 0.25·S_new_vel + 0.20·S_stale + 0.15·S_uc  (percentile-scored, 0–100)"),
        ("Liquidity Index",        "0.50·S_absorption + 0.30·S_fast_dom + 0.20·S_low_stale  (0–100)"),
        ("Demand Strength Index",  "0.35·S_absorption + 0.25·Liquidity + 0.20·Momentum + 0.20·CirclePremium  (0–100)"),
        ("Market Heat Index",      "50 + 0.5·Demand − 0.5·Supply  → clipped to [0, 100]"),
        ("Market Label",           "Hot (≥70) / Seller's (≥60) / Balanced / Buyer's (<40) / Cold (<30)"),
    ],
    hdr_color="244061"
)

add_heading(doc, "5.3 Buy Decision Engine", level=2)
doc.add_paragraph(
    "For any property with forecast data the system generates a Buy / Watch / Avoid "
    "recommendation backed by a quantitative score:"
)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Score  =  clip( 0.40·V + 0.35·G + 0.25·U − Risk_penalty,  0, 100 )")
run.bold = True
run.font.color.rgb = RGBColor(31, 78, 121)

doc.add_paragraph()
add_kv_table(doc, [
    ("V — Valuation Score",  "clip(50 − valuation_gap_pct × 2, 0, 100)  — reward for below-fair-value listings"),
    ("G — Growth Score",     "clip(50 + avg_YoY_pct × 3, 0, 100)  — reward for strong historical YoY"),
    ("U — Upside Score",     "clip(50 + expected_upside_pct × 1.5, 0, 100)  — reward for long-term forecast upside"),
    ("Risk Penalty",         "clip(QoQ_volatility_pct × 2, 0, 30)  — penalise erratic forecast trajectories"),
    ("Recommendation",       "Buy if Score ≥ 68  |  Watch if Score ≥ 50  |  Avoid otherwise"),
], hdr_color="1F4E79")

add_heading(doc, "5.4 ROI Calculator", level=2)
doc.add_paragraph(
    "A full discounted-cashflow style ROI analysis is computed for any holding period:"
)
add_bullet(doc, "Total Invested = buy_price × (1 + purchase_cost_pct)")
add_bullet(doc, "Gross Sale Price = forecast exit ppsf × area_sqft")
add_bullet(doc, "Net Sale Proceeds = Gross Sale × (1 − exit_cost_pct)")
add_bullet(doc, "Rental Income = monthly_rent (from dedicated rent ML model) × 12 × hold_years")
add_bullet(doc, "Holding Costs = buy_price × annual_holding_cost_pct × hold_years")
add_bullet(doc, "Net Profit = Net Sale + Rental − Holding − Total Invested")
add_bullet(doc, "ROI % = Net Profit / Total Invested × 100")
add_bullet(doc, "CAGR = (payoff_multiple)^(1/hold_years) − 1")

doc.add_paragraph()

# ── 6. Explainable AI ────────────────────────────────────────────────────────
add_heading(doc, "6. Explainable AI (XAI)", level=1, color=(31, 78, 121))
doc.add_paragraph(
    "Every prediction is accompanied by a human-readable explanation broken into 7 scored drivers "
    "(each −2 to +2). A natural-language narrative identifies the top 3 most influential factors."
)

add_metric_table(doc,
    ["Driver", "What It Measures", "Key Signal"],
    [
        ("Location / Market Zone",      "Circle rate vs. NCR segment median",           "cr_ratio = circle_rate / ncr_median_cr"),
        ("Circle-Rate Impact",          "Market premium over government floor value",    "pred_ratio = predicted_ppsf / circle_rate"),
        ("Built-Up / Plot Area",        "Size vs. NCR segment median area",             "area_ratio = area_sqft / ncr_median_area"),
        ("Road Connectivity",           "Proximity to NH/SH/MDR or main road",          "min(nh_km, sh_km, mdr_km) for plots; is_main_road for APT/BF"),
        ("Property-Type Quality",       "Segment tier, floor, furnishing, amenities",   "Composite score summed from individual flags"),
        ("Market Growth",               "Locality YoY forecast appreciation %",         "From rho-decay forecast engine"),
        ("Listing vs Segment Median",   "Predicted ppsf vs. NCR median ppsf",           "ppsf_ratio = pred_ppsf / ncr_median_ppsf"),
    ],
    hdr_color="244061"
)

# ── 7. Tech Stack ─────────────────────────────────────────────────────────────
add_heading(doc, "7. Technology Stack", level=1, color=(31, 78, 121))

add_metric_table(doc,
    ["Layer", "Technology", "Key Libraries / Tools"],
    [
        ("ML / Data Science",   "Python 3.14",          "scikit-learn, pandas, numpy, scipy, joblib, MLflow"),
        ("NLP Feature Eng.",    "Python",               "Regex-based NLP parser for BHK / area / furnishing extraction from listing descriptions"),
        ("Spatial Features",    "Python",               "KMeans (scikit-learn) for Voronoi/cluster spatial encoding; Shapely / GeoPandas for road distance"),
        ("API Backend",         "FastAPI (Python)",     "Pydantic, Uvicorn, concurrent.futures.ThreadPoolExecutor for parallel model loading"),
        ("Experiment Tracking", "MLflow",               "SQLite backend, per-segment experiments, versioned run metadata"),
        ("Frontend",            "React 18 + Vite",      "Recharts for all charts, Tailwind CSS for styling, single App.jsx component"),
        ("Containerisation",    "Docker",               "Dockerfile provided for backend; frontend served via Vite dev server or static build"),
        ("Data Storage",        "CSV / JSON",           "No database required for predictions; raw data in real_estate_data/, artifacts in artifact/ and opt/"),
    ],
    hdr_color="1F4E79"
)

# ── 8. Coverage ───────────────────────────────────────────────────────────────
add_heading(doc, "8. Geographic & Segment Coverage", level=1, color=(31, 78, 121))

add_metric_table(doc,
    ["Dimension", "Coverage"],
    [
        ("Cities",               "Delhi, Noida, Gurgaon, Faridabad, Ghaziabad, Greater Noida, Jaipur"),
        ("Property Types",       "Apartment, Builder Floor, Plot / Land"),
        ("Price Range (APT)",    "Budget (INR 2,000–6,000 /sqft) to Luxury (INR 20,000+ /sqft)"),
        ("Data Volume",          "~300,000 listing events across both scrapers"),
        ("Forecast Horizon",     "5 years, quarterly granularity (20 data points per property)"),
        ("Model Versions",       "APT v3 (1.85 GB), Builder Floor v2 (858 MB), Plot v1 (11.5 MB)"),
    ],
    hdr_color="244061"
)

# ── 9. Model Performance ──────────────────────────────────────────────────────
add_heading(doc, "9. Model Performance Summary", level=1, color=(31, 78, 121))
doc.add_paragraph(
    "Performance is reported in real price-per-sqft units (INR/sqft) so business "
    "stakeholders can interpret results without ML knowledge."
)

add_metric_table(doc,
    ["Segment", "R²", "MAPE", "Evaluation Units"],
    [
        ("Apartment",      "~0.85–0.90", "~8–12%",  "INR/sqft (back-calculated from log-ratio)"),
        ("Builder Floor",  "~0.83–0.88", "~9–13%",  "INR/sqft (back-calculated from log-ratio)"),
        ("Plot / Land",    "~0.80–0.86", "~10–15%", "INR/sqft (direct expm1 inversion)"),
    ],
    hdr_color="1F4E79"
)
doc.add_paragraph(
    "Note: Exact per-run metrics are stored in opt/{segment}/metrics.json and "
    "artifact/mlflow.db (MLflow tracking server)."
)

# ── 10. API Endpoints ─────────────────────────────────────────────────────────
add_heading(doc, "10. Key API Endpoints", level=1, color=(31, 78, 121))

add_metric_table(doc,
    ["Endpoint", "Method", "Description"],
    [
        ("/predict/apartment",             "POST", "Predict apartment price (ppsf + total) with XAI explanation"),
        ("/predict/builder-floor",         "POST", "Predict builder floor price with XAI explanation"),
        ("/predict/plot",                  "POST", "Predict plot price with road-connectivity XAI"),
        ("/meta/options",                  "GET",  "Return all dropdown options (cities, localities, BHK, age, furnishing, etc.)"),
        ("/meta/model-status",             "GET",  "Return loaded model health status"),
        ("/forecast/context",              "GET",  "Full forecast context for a property (quarter table, YoY series, rho payload, distribution)"),
        ("/forecast/buy-decision",         "GET",  "Buy/Watch/Avoid score and recommendation"),
        ("/forecast/roi",                  "GET",  "Full ROI calculation with rental income model"),
        ("/locality/intelligence",         "GET",  "Locality profile: ppsf, affordability, volatility, trend, forecast"),
        ("/market-intelligence/context",   "GET",  "Market heat index, supply/demand KPIs, index series for a locality"),
        ("/circle-rate",                   "GET",  "Lookup circle rate for a city × locality"),
        ("/road-distances",                "GET",  "NH / SH / MDR distances from lat/lon coordinates"),
    ],
    hdr_color="244061"
)

# ── 11. Frontend Dashboard ────────────────────────────────────────────────────
add_heading(doc, "11. Frontend Dashboard", level=1, color=(31, 78, 121))

doc.add_paragraph(
    "The React dashboard (frontend/src/App.jsx) provides six tab-based views, "
    "each consuming one or more API endpoints:"
)

add_metric_table(doc,
    ["Tab", "Functionality"],
    [
        ("Apartment Valuator",      "BHK / area / amenities form → instant price prediction + 7-driver XAI bar chart + narrative"),
        ("Builder Floor Valuator",  "Same form layout as apartment but with floor-level inputs → BF-specific XAI"),
        ("Plot / Land Valuator",    "Area / road width / usage type / lat-lon → plot prediction + highway-proximity XAI"),
        ("Forecast Intelligence",   "Locality search → 5-year forecast chart, YoY trend, distribution band, quarter table, buy/watch/avoid badge, ROI panel"),
        ("Locality Intelligence",   "Locality search → affordability score, listing count, price trend, volatility gauge, forecast appreciation"),
        ("Market Intelligence",     "Locality search → market heat index gauge, supply/demand KPI cards, absorption chart, price-cut/hike trend"),
    ],
    hdr_color="1F4E79"
)

# ── 12. Design Decisions ──────────────────────────────────────────────────────
add_heading(doc, "12. Key Design Decisions & Rationale", level=1, color=(31, 78, 121))

add_bullet(doc,
    "Predicting price_per_sqft / circle_rate instead of raw price normalises across cities and price tiers, "
    "significantly improving model generalisation.",
    "Circle-Rate Normalisation —"
)
add_bullet(doc,
    "KMeans Voronoi clusters on lat/lon capture spatial autocorrelation without requiring a fixed "
    "grid or administrative boundary definitions.",
    "Spatial Encoding —"
)
add_bullet(doc,
    "The rho-decay model provides an interpretable blend between short-term market price and "
    "long-term fundamental value — a design inspired by Bayesian updating.",
    "Rho-Decay Forecast —"
)
add_bullet(doc,
    "All models load concurrently at startup using ThreadPoolExecutor, keeping cold-start "
    "latency under 30 seconds despite 3 GB of total model weights.",
    "Parallel Model Loading —"
)
add_bullet(doc,
    "A dedicated rent model per segment (apt_rent_model, bf_rent_model) enables "
    "yield-based ROI calculations without hardcoded assumptions.",
    "Separate Rent Models —"
)
add_bullet(doc,
    "All indices (supply pressure, liquidity, market heat) use within-city percentile "
    "scoring so cities with structurally different price levels remain comparable.",
    "Percentile-Scored Indices —"
)

doc.add_paragraph()

# ── 13. Glossary ──────────────────────────────────────────────────────────────
add_heading(doc, "13. Glossary", level=1, color=(31, 78, 121))

add_kv_table(doc, [
    ("Circle Rate",       "Government-mandated minimum property valuation floor (INR/sqft) set by state authorities"),
    ("PPSF",             "Price Per Square Foot — primary unit for all valuations in this system"),
    ("Rho (ρ)",          "Blend weight in forecast model controlling how much the listing price vs. ML model price drives future value"),
    ("I_l,t",            "Locality growth index at quarter t — derived from GSDP, RBI repo rate, and home loan rate series"),
    ("Voronoi Features", "Spatial cluster dummy features assigned by KMeans clustering of lat/lon coordinates"),
    ("Absorption Rate",  "Fraction of active listings that disappear between consecutive scrape dates — proxy for buyer demand"),
    ("Market Heat Index","Composite 0–100 index: >70 = Hot market; <30 = Cold market"),
    ("XAI",             "Explainable AI — 7 scored drivers that decompose why a property is predicted at a specific price"),
    ("YoY",             "Year-over-Year price appreciation percentage"),
    ("CAGR",            "Compound Annual Growth Rate — annualised investment return metric"),
    ("MLflow",          "Open-source ML experiment tracking library — stores model runs, metrics, and parameters"),
    ("NCR",             "National Capital Region — metropolitan region including Delhi, Noida, Gurgaon, Faridabad, Ghaziabad, Greater Noida"),
], hdr_color="1F4E79")

# ── Footer note ───────────────────────────────────────────────────────────────
doc.add_page_break()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("GeoAI Real-Estate Intelligence Platform  —  Confidential Technical Reference")
run.italic = True
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(150, 150, 150)
p.add_run(f"\nGenerated automatically on {datetime.date.today().strftime('%B %d, %Y')}").italic = True

# ── Save ──────────────────────────────────────────────────────────────────────
out_path = "GeoAI_RealEstate_System_Overview.docx"
doc.save(out_path)
print(f"Document saved: {out_path}")
