"""
Generate a Microsoft Word document explaining all Decision Intelligence formulas
and a fully worked example using a 3BHK apartment in Noida Extension.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ──────────────────────────────────────────────────────────────
# Helpers
# ──────────────────────────────────────────────────────────────

BRAND_BLUE   = RGBColor(0x1A, 0x37, 0x6C)   # dark navy
ACCENT_BLUE  = RGBColor(0x27, 0x6E, 0xF7)   # vivid blue
SECTION_BG   = RGBColor(0xEF, 0xF4, 0xFF)   # light blue tint
GREEN        = RGBColor(0x1A, 0x7A, 0x3C)
RED          = RGBColor(0xB3, 0x0C, 0x0C)
GOLD         = RGBColor(0xB8, 0x86, 0x0B)
DARK_GREY    = RGBColor(0x33, 0x33, 0x33)
MID_GREY     = RGBColor(0x66, 0x66, 0x66)
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)
TABLE_HEADER = RGBColor(0x1A, 0x37, 0x6C)
TABLE_ALT    = RGBColor(0xF2, 0xF6, 0xFF)


def _rgb_hex(rgb) -> str:
    # RGBColor is a tuple subclass; str(rgb) returns the hex string directly
    return str(rgb).upper()


def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    hex_color = _rgb_hex(rgb)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_borders(table):
    """Thin borders on every cell."""
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
                border = OxmlElement(f"w:{side}")
                border.set(qn("w:val"), "single")
                border.set(qn("w:sz"), "4")
                border.set(qn("w:space"), "0")
                border.set(qn("w:color"), "C0CADE")
                tcBorders.append(border)
            tcPr.append(tcBorders)


def header_row(table, *texts, bg=TABLE_HEADER):
    row = table.rows[0]
    for i, txt in enumerate(texts):
        cell = row.cells[i]
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(txt)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(9)


def data_row(table, row_idx, *texts, bold_first=False, alt=False, aligns=None):
    row = table.rows[row_idx]
    for i, txt in enumerate(texts):
        cell = row.cells[i]
        if alt:
            set_cell_bg(cell, TABLE_ALT)
        p = cell.paragraphs[0]
        align = (aligns[i] if aligns and i < len(aligns) else WD_ALIGN_PARAGRAPH.LEFT)
        p.alignment = align
        run = p.add_run(str(txt))
        run.font.size = Pt(9)
        if bold_first and i == 0:
            run.bold = True
        run.font.color.rgb = DARK_GREY


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    if level == 0:
        run.font.size = Pt(20)
        run.font.color.rgb = BRAND_BLUE
    elif level == 1:
        run.font.size = Pt(14)
        run.font.color.rgb = BRAND_BLUE
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after  = Pt(4)
        # bottom border
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "1A376C")
        pBdr.append(bottom)
        pPr.append(pBdr)
    elif level == 2:
        run.font.size = Pt(11)
        run.font.color.rgb = ACCENT_BLUE
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(2)
    elif level == 3:
        run.font.size = Pt(10)
        run.font.color.rgb = DARK_GREY
        p.paragraph_format.space_before = Pt(6)
    return p


def add_body(doc, text, color=None, italic=False, size=9.5):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.italic = italic
    run.font.color.rgb = color or DARK_GREY
    p.paragraph_format.space_after = Pt(3)
    return p


def add_formula_block(doc, formula_text):
    """Monospaced, indented formula paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(formula_text)
    run.font.name = "Courier New"
    run.font.size = Pt(9.5)
    run.font.color.rgb = ACCENT_BLUE
    run.bold = True
    return p


def add_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run("⚠  " + text)
    run.font.size = Pt(8.5)
    run.italic = True
    run.font.color.rgb = GOLD
    return p


def add_result_box(doc, text, color=GREEN):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run("▶  " + text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = color
    return p


def add_page_break(doc):
    doc.add_page_break()


# ──────────────────────────────────────────────────────────────
# Build Document
# ──────────────────────────────────────────────────────────────

def build():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Cover ──────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("GeoAI Real Estate Platform")
    r.font.size = Pt(22)
    r.bold = True
    r.font.color.rgb = BRAND_BLUE

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Decision Intelligence — Complete Formula Reference")
    r2.font.size = Pt(14)
    r2.font.color.rgb = ACCENT_BLUE
    r2.bold = True

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("With Fully Worked Example · 3 BHK Apartment · Noida Extension")
    r3.font.size = Pt(11)
    r3.font.color.rgb = MID_GREY
    r3.italic = True

    doc.add_paragraph()
    add_body(doc,
        "This document covers every formula used by the platform's three intelligence layers: "
        "(1) Market Intelligence Indices, (2) Buy / Watch / Avoid Decision Score, and "
        "(3) ROI & Cashflow Projection. A single 3 BHK apartment example in Noida Extension "
        "is traced through all three layers end-to-end.",
        size=9.5, italic=True, color=MID_GREY)

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════
    # PART A — MARKET INTELLIGENCE INDICES
    # ══════════════════════════════════════════════════════════
    add_heading(doc, "PART A — Market Intelligence Indices", level=0)
    add_body(doc,
        "Five market indices and one composite heat score are computed for every "
        "(segment, city, locality) group from raw listing events. All percentile-based "
        "scores are ranked within the city peer group so Noida Extension is compared "
        "only to other Greater Noida localities.")

    # ── A1. Percentile Scoring ─────────────────────────────────
    add_heading(doc, "A1. Percentile Scoring — The Foundation", level=1)
    add_body(doc,
        "Metrics that are 'higher = worse for buyers' (e.g. active supply) or "
        "'lower = better' (e.g. days on market) are first ranked by percentile "
        "within the city, then optionally inverted. This ensures every metric "
        "is comparable on the same 0–100 scale regardless of its absolute units.")

    add_formula_block(doc,
        "percentile_score(x, higher_is_better=True)\n"
        "  = rank_pct(x within city) × 100           if higher_is_better\n"
        "  = 100 - rank_pct(x within city) × 100     if lower_is_better")

    add_body(doc,
        "rank_pct is pandas .rank(pct=True), so a value at the 70th percentile "
        "returns 70.0. This makes the score interpretation universal: "
        "70 always means 'better than 70 % of peers'.")

    # ── A2. Supply Pressure Index ──────────────────────────────
    add_heading(doc, "A2. Supply Pressure Index  (0 – 100, higher = more supply pressure)", level=1)
    add_body(doc,
        "Measures how much new inventory is entering the market relative to peers. "
        "A high score is unfavourable for buyers — supply is outpacing demand.")

    t = doc.add_table(rows=5, cols=3)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    header_row(t, "Component", "Raw Metric", "Weight")
    data_row(t, 1, "S_active",  "Active listing stock (percentile, higher = more supply)", "40 %", aligns=[WD_ALIGN_PARAGRAPH.LEFT]*3)
    data_row(t, 2, "S_new_vel", "New-listing velocity = new_count / active_stock (percentile)", "25 %", alt=True)
    data_row(t, 3, "S_stale",   "Stale share >90 days DOM (percentile, higher = more stale)", "20 %")
    data_row(t, 4, "S_uc",      "Under-construction share (percentile)", "15 %", alt=True)
    set_cell_borders(t)

    doc.add_paragraph()
    add_formula_block(doc,
        "Supply_Pressure = clamp(\n"
        "    0.40 × S_active  +  0.25 × S_new_vel\n"
        "  + 0.20 × S_stale   +  0.15 × S_uc ,  0, 100)")

    # ── A3. Liquidity Index ────────────────────────────────────
    add_heading(doc, "A3. Liquidity Index  (0 – 100, higher = more liquid / easier to exit)", level=1)
    add_body(doc,
        "Measures how quickly properties trade hands. "
        "Absorption rate (proxy for sell-throughs between scrape cycles) "
        "carries the most weight; fast DOM and low stale inventory reinforce it.")

    add_formula_block(doc,
        "S_absorption = percentile(absorption_rate,  higher_is_better=True )\n"
        "S_fast_dom   = percentile(median_dom,        higher_is_better=False)  ← low DOM is good\n"
        "S_low_stale  = percentile(stale_share,       higher_is_better=False)  ← low stale is good\n\n"
        "Liquidity = clamp(\n"
        "    0.50 × S_absorption  +  0.30 × S_fast_dom  +  0.20 × S_low_stale ,  0, 100)")

    add_body(doc,
        "absorption_rate = (listings present in current scrape but absent in next scrape) "
        "/ active_supply — computed only when the next scrape also visited the same locality.")

    # ── A4. Price Momentum Score ───────────────────────────────
    add_heading(doc, "A4. Price Momentum Score  (0 – 100, 50 = neutral)", level=1)
    add_body(doc,
        "Tracks whether sellers are raising or cutting prices on re-listed properties. "
        "Unlike other indices this is NOT percentile-based — it is a direct arithmetic score.")

    add_formula_block(doc,
        "f_hike  = price_hike_count  / updated_count   (fraction of updated listings with price up)\n"
        "f_cut   = price_cut_count   / updated_count   (fraction of updated listings with price down)\n\n"
        "Momentum = clamp(50 + 40 × f_hike − 40 × f_cut ,  0, 100)\n\n"
        "If updated_count = 0  →  Momentum = 50  (neutral, no data)")

    add_note(doc, "Equal hikers and cutters → score stays exactly 50. "
             "Dominantly rising prices → score > 50. Dominantly falling → score < 50.")

    # ── A5. Circle Premium Score ────────────────────────────────
    add_heading(doc, "A5. Circle Premium Score  (0 – 100)", level=1)
    add_body(doc,
        "Compares the market-observed median PPSF to the government circle rate. "
        "A ratio well above 1.0 means real demand is far higher than the regulatory floor — "
        "a demand confidence signal. Not percentile-based.")

    add_formula_block(doc,
        "relative = median_ppsf / circle_rate      (if circle_rate > 0)\n"
        "         = median_ppsf / city_median_ppsf  (fallback when circle rate unavailable)\n\n"
        "Circle_Premium = clamp(50 + 50 × (relative − 1),  0, 100)\n\n"
        "Examples:\n"
        "  relative = 1.0  →  50  (market trades exactly at govt floor)\n"
        "  relative = 2.0  →  100 (market is 2× govt floor; clipped at 100)\n"
        "  relative = 0.5  →  0   (market below govt floor; clipped at 0)")

    # ── A6. Demand Strength Index ──────────────────────────────
    add_heading(doc, "A6. Demand Strength Index  (0 – 100)", level=1)
    add_body(doc,
        "A composite demand signal built from absorption, liquidity, momentum and "
        "circle premium. Absorption carries the highest weight because it is the "
        "most direct signal that buyers are committing.")

    add_formula_block(doc,
        "Demand_Strength = clamp(\n"
        "    0.35 × S_absorption   (percentile of absorption rate)\n"
        "  + 0.25 × Liquidity      (the full Liquidity Index above)\n"
        "  + 0.20 × Momentum       (Price Momentum Score)\n"
        "  + 0.20 × Circle_Premium (Circle Premium Score)\n"
        ",  0, 100)")

    # ── A7. Market Heat Index ──────────────────────────────────
    add_heading(doc, "A7. Market Heat Index  (0 – 100) and Label", level=1)
    add_body(doc,
        "The single top-level market temperature score. High demand AND low supply pressure "
        "produces a hot market. Smoothed over a rolling 4-scrape window.")

    add_formula_block(doc,
        "Market_Heat = clamp(50 + 0.5 × Demand_Strength − 0.5 × Supply_Pressure,  0, 100)\n\n"
        "mhi_rolling4 = rolling mean of the last 4 scrape dates per locality")

    t2 = doc.add_table(rows=6, cols=2)
    t2.style = "Table Grid"
    t2.alignment = WD_TABLE_ALIGNMENT.LEFT
    header_row(t2, "Score Range", "Market Label")
    labels = [
        ("75 – 100", "Hot demand-led market"),
        ("60 – 74",  "Positive market"),
        ("45 – 59",  "Balanced market"),
        ("30 – 44",  "Supply-heavy market"),
        (" 0 – 29",  "Weak / stale market"),
    ]
    for i, (rng, lbl) in enumerate(labels, start=1):
        data_row(t2, i, rng, lbl, alt=(i % 2 == 0))
    set_cell_borders(t2)

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════
    # PART B — BUY / WATCH / AVOID SCORE
    # ══════════════════════════════════════════════════════════
    add_heading(doc, "PART B — Buy / Watch / Avoid Decision Score", level=0)
    add_body(doc,
        "Given a specific property listing, the platform runs four sub-scores and combines "
        "them into a single 0–100 investment recommendation score. Inputs come from "
        "(a) the ML valuation model, (b) the price forecast trajectory, and "
        "(c) historical YoY series.")

    # ── B1. Inputs ─────────────────────────────────────────────
    add_heading(doc, "B1. Key Inputs", level=1)

    t3 = doc.add_table(rows=7, cols=3)
    t3.style = "Table Grid"
    header_row(t3, "Variable", "Description", "Source")
    inputs_data = [
        ("listingPricePpsf",   "Asking price per sqft from the listing",            "User / listing data"),
        ("modelPricePpsf",     "ML model's fair-value estimate per sqft",            "Trained XGBoost model"),
        ("valuationGapPct (Δ%)", "(listingPpsf − modelPpsf) / modelPpsf × 100",       "Computed"),
        ("avgYoY%",            "Mean YoY% from the property's forecast series",      "Forecast model"),
        ("expectedUpside%",    "(forecastPpsf_horizon / listingPpsf − 1) × 100",     "Computed"),
        ("σ_QoQ (volatility)", "Std-dev of quarter-over-quarter % in forecast series","Computed"),
    ]
    for i, (v, d, s) in enumerate(inputs_data, start=1):
        data_row(t3, i, v, d, s, alt=(i % 2 == 0), bold_first=True)
    set_cell_borders(t3)

    # ── B2. Horizon PPSF ──────────────────────────────────────
    add_heading(doc, "B2. Horizon Price Extraction", level=1)
    add_body(doc,
        "The forecast model produces quarterly price-per-sqft predictions. "
        "The exit price at the investor's chosen hold period is read off the series.")

    add_formula_block(doc,
        "horizon_quarter = hold_years × 4      (e.g. 5 years → Q20)\n"
        "exitPpsf        = forecast_series[horizon_quarter − 1]\n\n"
        "expectedUpside% = (exitPpsf / listingPpsf − 1) × 100")

    # ── B3. Sub-scores ─────────────────────────────────────────
    add_heading(doc, "B3. Four Sub-Scores  (each 0 – 100)", level=1)
    add_body(doc, "Each factor is mapped to a 0–100 scale and then combined.")

    add_formula_block(doc,
        "ValuationScore = clamp(50 − Δ% × 2 ,  0, 100)\n"
        "  Δ% > 0  → listing overpriced → score < 50 (negative signal)\n"
        "  Δ% < 0  → listing underpriced → score > 50 (positive signal)\n\n"
        "GrowthScore    = clamp(50 + avgYoY% × 3,  0, 100)\n"
        "  Each 1% of annual growth adds 3 points above the neutral 50\n\n"
        "UpsideScore    = clamp(50 + expectedUpside% × 1.5,  0, 100)\n"
        "  50% total upside over the horizon → score = 125 → clipped to 100\n\n"
        "RiskPenalty    = clamp(σ_QoQ × 2,  0, 30)\n"
        "  High forecast volatility deducts up to 30 points from the final score")

    # ── B4. Overall Score ──────────────────────────────────────
    add_heading(doc, "B4. Overall Score and Recommendation", level=1)

    add_formula_block(doc,
        "OverallScore = clamp(\n"
        "    0.40 × ValuationScore\n"
        "  + 0.35 × GrowthScore\n"
        "  + 0.25 × UpsideScore\n"
        "  − RiskPenalty\n"
        ",  0, 100)")

    t4 = doc.add_table(rows=4, cols=3)
    t4.style = "Table Grid"
    header_row(t4, "Score Range", "Recommendation", "Meaning")
    recs = [
        ("68 – 100", "BUY",   "Strong investment case — valuation, growth and upside aligned"),
        ("50 – 67",  "WATCH", "Reasonable case but one or more signals are not fully aligned"),
        (" 0 – 49",  "AVOID", "Weak fundamentals — overpriced, low growth or high volatility"),
    ]
    for i, (rng, rec, meaning) in enumerate(recs, start=1):
        data_row(t4, i, rng, rec, meaning, alt=(i % 2 == 0))
    set_cell_borders(t4)

    # ── B5. Confidence ─────────────────────────────────────────
    add_heading(doc, "B5. Confidence Score  (0.35 – 0.95)", level=1)
    add_body(doc, "Reflects data quality, not investment quality.")
    add_formula_block(doc,
        "confidence = 0.55\n"
        "           + min(forecast_quarters_available, 20) / 20 × 0.25   ← data depth\n"
        "           + 0.10   if locality-level forecast (not city fallback)\n"
        "           − 0.10   if city-level fallback used\n"
        "= clamp(result,  0.35, 0.95)")

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════
    # PART C — ROI & CASHFLOW MODEL
    # ══════════════════════════════════════════════════════════
    add_heading(doc, "PART C — ROI & Cashflow Projection Model", level=0)
    add_body(doc,
        "Converts the forecast exit price, a user-specified holding period, cost "
        "assumptions, and either a user-supplied or model-predicted rent yield into a "
        "simplified cashflow P&L and annualised return metrics.")

    # ── C1. Inputs ─────────────────────────────────────────────
    add_heading(doc, "C1. Input Parameters", level=1)

    t5 = doc.add_table(rows=9, cols=3)
    t5.style = "Table Grid"
    header_row(t5, "Parameter", "Default / Typical", "Description")
    roi_inputs = [
        ("hold_years",             "5 years",   "Investor's intended holding period"),
        ("area_sqft",              "Property",  "Built-up area in sq ft; falls back to 1,000 sqft if missing"),
        ("purchaseCostPct",        "7 %",       "Stamp duty + registration + misc acquisition costs"),
        ("annualHoldingCostPct",   "1 %",       "Property tax, maintenance, insurance per year"),
        ("exitCostPct",            "2 %",       "Brokerage + legal fees at time of sale"),
        ("rentYieldPct",           "Model / User", "Annual rent as % of purchase price; predicted by rent ML model if not supplied"),
        ("buyPpsf",                "Listing PPSF", "Entry price per sqft; falls back to model PPSF if listing unavailable"),
        ("exitPpsf",               "Forecast Q(hold_years × 4)", "ML-forecast exit price per sqft at the horizon quarter"),
    ]
    for i, (p, d, desc) in enumerate(roi_inputs, start=1):
        data_row(t5, i, p, d, desc, alt=(i % 2 == 0), bold_first=True)
    set_cell_borders(t5)

    # ── C2. Cashflow Formulas ──────────────────────────────────
    add_heading(doc, "C2. Cashflow Calculation — Step by Step", level=1)

    steps = [
        ("Step 1 — Entry Cost",
         "buyPrice         = buyPpsf × area_sqft\n"
         "purchaseCosts    = buyPrice × (purchaseCostPct / 100)\n"
         "totalInvested    = buyPrice + purchaseCosts"),
        ("Step 2 — Exit Proceeds",
         "grossSalePrice   = exitPpsf × area_sqft\n"
         "exitCosts        = grossSalePrice × (exitCostPct / 100)\n"
         "netSaleProceeds  = grossSalePrice − exitCosts"),
        ("Step 3 — Rental Income",
         "annualRent       = buyPrice × (rentYieldPct / 100)\n"
         "rentalIncomeTotal = annualRent × hold_years"),
        ("Step 4 — Holding Costs",
         "holdingCostsTotal = buyPrice × (annualHoldingCostPct / 100) × hold_years"),
        ("Step 5 — Net Profit",
         "netProfit = netSaleProceeds + rentalIncomeTotal\n"
         "          − holdingCostsTotal − totalInvested"),
    ]
    for title, formula in steps:
        add_heading(doc, title, level=3)
        add_formula_block(doc, formula)

    # ── C3. Return Metrics ─────────────────────────────────────
    add_heading(doc, "C3. Return Metrics", level=1)

    add_formula_block(doc,
        "ROI %           = (netProfit / totalInvested) × 100\n\n"
        "payoffMultiple  = (netSaleProceeds + rentalIncomeTotal − holdingCostsTotal)\n"
        "                  / totalInvested\n\n"
        "CAGR %          = (payoffMultiple ^ (1 / hold_years) − 1) × 100")

    t6 = doc.add_table(rows=4, cols=2)
    t6.style = "Table Grid"
    header_row(t6, "ROI %", "Verdict")
    verdicts = [("≥ 35 %", "STRONG — Excellent capital return"), ("15 – 34 %", "MODERATE — Acceptable return"), ("< 15 %", "WEAK — Below par")]
    for i, (rng, v) in enumerate(verdicts, start=1):
        data_row(t6, i, rng, v, alt=(i % 2 == 0))
    set_cell_borders(t6)

    # ── C4. Rent Model ─────────────────────────────────────────
    add_heading(doc, "C4. Rent Model (when rentYieldPct is not supplied by user)", level=1)
    add_body(doc,
        "For Apartment and Builder Floor segments, the platform calls a dedicated "
        "XGBoost rent model (trained on log-transformed monthly rent) to estimate yield. "
        "Plots default to zero rental income as a conservative assumption.")
    add_formula_block(doc,
        "log_monthly_rent = rent_model.predict(property_features)\n"
        "monthly_rent     = exp(log_monthly_rent) − 1      ← np.expm1 inverse of log1p\n"
        "annual_rent      = monthly_rent × 12\n"
        "effectiveYield%  = (annual_rent / buyPrice) × 100")

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════
    # PART D — WORKED EXAMPLE
    # ══════════════════════════════════════════════════════════
    add_heading(doc, "PART D — Fully Worked Example: 3 BHK Apartment, Noida Extension", level=0)
    add_body(doc,
        "All numbers below are derived from actual platform data (market intelligence "
        "artifact + Noida/Alpha-1 zone forecast as the nearest comparable forecast zone "
        "because Noida Extension is classified under Greater Noida in our locality mapping).")

    # ── D1. Property Inputs ────────────────────────────────────
    add_heading(doc, "D1. Property Inputs", level=1)

    t7 = doc.add_table(rows=16, cols=3)
    t7.style = "Table Grid"
    header_row(t7, "Input Field", "Value", "Notes")
    prop_inputs = [
        ("Segment",             "Apartment",            "Flat in a multi-storey society"),
        ("Locality",            "Noida Extension",      "Greater Noida West, Greater Noida"),
        ("BHK",                 "3 BHK",                "3 bedrooms"),
        ("Covered Area",        "1,500 sqft",           "—"),
        ("Bathrooms",           "3",                    "—"),
        ("Balconies",           "2",                    "—"),
        ("Floor Level",         "Mid (5th of 15)",      "floor_medium = 1"),
        ("Age",                 "New Construction",     "< 5 years"),
        ("Furnishing",          "Semi-Furnished",       "—"),
        ("Facing",              "North",                "—"),
        ("Is Parking",          "Yes (1)",              "—"),
        ("Is Gated Society",    "Yes (1)",              "—"),
        ("Listing Price",       "₹9,500 / sqft",        "₹1,42,50,000 total"),
        ("ML Model Fair Value", "₹8,862 / sqft",        "From XGBoost model; ≈ locality median"),
        ("Circle Rate",         "₹3,902 / sqft",        "Government floor, Greater Noida Residential"),
    ]
    for i, (f, v, n) in enumerate(prop_inputs, start=1):
        data_row(t7, i, f, v, n, alt=(i % 2 == 0), bold_first=True)
    set_cell_borders(t7)

    # ── D2. Market Intelligence ────────────────────────────────
    add_heading(doc, "D2. Market Intelligence Scores — Noida Extension (Actual Data)", level=1)

    t8 = doc.add_table(rows=13, cols=3)
    t8.style = "Table Grid"
    header_row(t8, "Metric", "Raw Value", "Score / Comment")
    mi_data = [
        ("Active Supply Stock",         "905 listings",    "88th percentile in Greater Noida — very high"),
        ("New Supply Velocity",          "93.15 %",        "95th percentile — massive fresh inventory"),
        ("Stale Inventory Share",        "2.1 %",          "12th percentile — very little old stock"),
        ("Under-Construction Share",     "12.49 %",        "55th percentile — moderate pipeline"),
        ("Supply Pressure Index",        "—",              "= 0.40(88) + 0.25(95) + 0.20(12) + 0.15(55) = 71.4"),
        ("Absorption Rate",              "0.0 %",          "~10th percentile — no tracked exits"),
        ("Median Days on Market",        "4 days",         "~90th percentile inverted — fast churn"),
        ("Liquidity Index",              "—",              "= 0.50(10) + 0.30(90) + 0.20(88) = 41.3"),
        ("Price Hike / Cut Frequency",   "1.61 % each",    "Equal → Momentum = 50 + 40(0.0161) − 40(0.0161) = 50.0"),
        ("Median PPSF / Circle Rate",    "₹8,862 / ₹3,902","relative = 2.271 → Circle Premium = 50 + 50×1.271 = 113.55 → clipped to 100.0"),
        ("Demand Strength Index",        "—",              "= 0.35(35) + 0.25(41.3) + 0.20(50) + 0.20(100) = 52.5"),
        ("Market Heat Index / Label",    "—",              "= 50 + 0.5(52.5) − 0.5(71.4) = 40.6  →  'Supply-heavy market'"),
    ]
    for i, (m, rv, sc) in enumerate(mi_data, start=1):
        data_row(t8, i, m, rv, sc, alt=(i % 2 == 0), bold_first=True)
    set_cell_borders(t8)

    add_result_box(doc,
        "Market verdict: Supply-heavy (40.6). 905 new listings are flooding in, but market "
        "prices still trade at 2.27× the government circle rate — underlying demand is real.",
        color=GOLD)

    # ── D3. Buy/Watch/Avoid ────────────────────────────────────
    add_heading(doc, "D3. Buy / Watch / Avoid Score — Step-by-Step", level=1)

    add_heading(doc, "Forecast Series (Noida/Alpha-1 zone, quarterly)", level=2)
    add_body(doc, "The forecast model covers 20 quarters (5 years) from Q1 2026 onwards:")

    t9 = doc.add_table(rows=6, cols=5)
    t9.style = "Table Grid"
    header_row(t9, "Quarter", "Date", "Forecast PPSF (₹)", "QoQ %", "Note")
    forecast_pts = [
        ("Q1",  "Jul 2026",  "9,943",  "—",    "Entry baseline"),
        ("Q4",  "Apr 2027",  "10,967", "+10.3 % vs Q1", "Year 1 end"),
        ("Q8",  "Apr 2028",  "12,978", "+18.3 % vs Q4", "Year 2 end"),
        ("Q12", "Apr 2029",  "14,703", "+13.3 % vs Q8", "Year 3 end"),
        ("Q20", "Apr 2031",  "16,695", "+13.6 % vs Q16","Year 5 end — exit price used"),
    ]
    for i, (q, dt, ppsf, qoq, note) in enumerate(forecast_pts, start=1):
        data_row(t9, i, q, dt, ppsf, qoq, note, alt=(i % 2 == 0))
    set_cell_borders(t9)

    add_heading(doc, "Inputs to the Decision Score", level=2)

    add_formula_block(doc,
        "listingPpsf         = ₹9,500 / sqft\n"
        "modelPpsf           = ₹8,862 / sqft\n"
        "Δ% (valuationGap)   = (9,500 − 8,862) / 8,862 × 100 = +7.2 %  (overpriced by 7.2 %)\n\n"
        "exitPpsf (Q20)      = ₹16,695 / sqft\n"
        "expectedUpside%     = (16,695 / 9,500 − 1) × 100 = +75.7 %\n\n"
        "avgYoY%             ≈ 11.1 %  (mean of 5 annual steps in the forecast)\n"
        "σ_QoQ (volatility)  ≈ 1.5 %  (std-dev of quarterly % changes in forecast series)")

    add_heading(doc, "Sub-Score Calculations", level=2)

    add_formula_block(doc,
        "ValuationScore = clamp(50 − 7.2 × 2 ,  0, 100)  =  35.6\n"
        "GrowthScore    = clamp(50 + 11.1 × 3,  0, 100)  =  83.3\n"
        "UpsideScore    = clamp(50 + 75.7 × 1.5, 0, 100) = 163.6 → clipped to 100.0\n"
        "RiskPenalty    = clamp(1.5 × 2,  0, 30)         =   3.0")

    add_heading(doc, "Overall Score", level=2)

    add_formula_block(doc,
        "OverallScore = 0.40 × 35.6  +  0.35 × 83.3  +  0.25 × 100.0  −  3.0\n"
        "            =  14.24        +   29.16        +   25.0          −  3.0\n"
        "            =  65.4")

    add_result_box(doc,
        "Score 65.4  →  WATCH (threshold: Buy ≥ 68, Watch ≥ 50, Avoid < 50)",
        color=GOLD)

    add_heading(doc, "Sensitivity: What if you negotiate to fair value?", level=2)
    add_formula_block(doc,
        "If listing = model fair value  →  Δ% = 0\n"
        "ValuationScore = clamp(50 − 0 × 2, 0, 100) = 50.0\n"
        "OverallScore   = 0.40×50 + 0.35×83.3 + 0.25×100 − 3.0  =  71.2  →  BUY")

    add_result_box(doc,
        "Negotiating from ₹9,500 to ₹8,862 / sqft flips the recommendation from WATCH to BUY.",
        color=GREEN)

    # ── D4. ROI ────────────────────────────────────────────────
    add_heading(doc, "D4. ROI & Cashflow — Full Calculation (5-Year Hold)", level=1)

    add_heading(doc, "Assumptions", level=2)

    t10 = doc.add_table(rows=7, cols=2)
    t10.style = "Table Grid"
    header_row(t10, "Parameter", "Value Used")
    roi_assm = [
        ("Hold period",            "5 years"),
        ("Area",                   "1,500 sqft"),
        ("Purchase costs",         "7 % of buy price  (stamp duty + registration)"),
        ("Annual holding costs",   "1 % of buy price per year  (maintenance + tax)"),
        ("Exit costs",             "2 % of gross sale price  (brokerage)"),
        ("Rent yield",             "3 % p.a. of buy price  (ML rent model estimate for 3BHK semi-furnished)"),
    ]
    for i, (p, v) in enumerate(roi_assm, start=1):
        data_row(t10, i, p, v, alt=(i % 2 == 0), bold_first=True)
    set_cell_borders(t10)

    add_heading(doc, "Step-by-Step Numbers", level=2)

    t11 = doc.add_table(rows=15, cols=3)
    t11.style = "Table Grid"
    header_row(t11, "Line Item", "Calculation", "Amount (₹)")
    cashflow_rows = [
        ("Buy Price",             "9,500 × 1,500",                          "1,42,50,000"),
        ("Purchase Costs (7%)",   "1,42,50,000 × 7 %",                      "    9,97,500"),
        ("TOTAL INVESTED",        "Buy Price + Purchase Costs",              "1,52,47,500"),
        ("",                      "",                                        ""),
        ("Gross Sale Price",      "16,695 × 1,500",                         "2,50,42,500"),
        ("Exit Costs (2%)",       "2,50,42,500 × 2 %",                      "    5,00,850"),
        ("NET SALE PROCEEDS",     "Gross Sale − Exit Costs",                 "2,45,41,650"),
        ("",                      "",                                        ""),
        ("Annual Rent",           "1,42,50,000 × 3 %",                      "    4,27,500"),
        ("RENTAL INCOME (5 yr)",  "4,27,500 × 5",                           "   21,37,500"),
        ("",                      "",                                        ""),
        ("HOLDING COSTS (5 yr)",  "1,42,50,000 × 1 % × 5",                 "    7,12,500"),
        ("",                      "",                                        ""),
        ("NET PROFIT",
         "Net Sale + Rental Income − Holding Costs − Total Invested",
         "1,07,19,150"),
    ]
    for i, (lbl, calc, amt) in enumerate(cashflow_rows, start=1):
        bold = lbl in ("TOTAL INVESTED", "NET SALE PROCEEDS", "RENTAL INCOME (5 yr)",
                       "HOLDING COSTS (5 yr)", "NET PROFIT")
        row = t11.rows[i]
        for j, txt in enumerate((lbl, calc, amt)):
            cell = row.cells[j]
            if lbl == "NET PROFIT":
                set_cell_bg(cell, RGBColor(0xE6, 0xF4, 0xEA))
            elif bold:
                set_cell_bg(cell, RGBColor(0xEE, 0xF3, 0xFF))
            elif not txt:
                set_cell_bg(cell, RGBColor(0xF8, 0xF8, 0xF8))
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if j == 2 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(txt)
            run.font.size = Pt(9)
            run.bold = bold
            run.font.color.rgb = GREEN if lbl == "NET PROFIT" else DARK_GREY
        set_cell_borders(t11)

    add_heading(doc, "Return Metrics", level=2)

    add_formula_block(doc,
        "ROI %          = (1,07,19,150 / 1,52,47,500) × 100          =  70.3 %\n\n"
        "payoffMultiple = (2,45,41,650 + 21,37,500 − 7,12,500)\n"
        "                  / 1,52,47,500\n"
        "               = 2,59,66,650 / 1,52,47,500                  =  1.703\n\n"
        "CAGR %         = (1.703 ^ (1/5) − 1) × 100\n"
        "               = (1.1124 − 1) × 100                         =  11.2 % per year")

    add_result_box(doc,
        "ROI = 70.3 %  →  STRONG  |  CAGR = 11.2 % p.a.  |  Payoff Multiple = 1.70×",
        color=GREEN)

    # ── D5. Summary ────────────────────────────────────────────
    add_heading(doc, "D5. Decision Summary", level=1)

    t12 = doc.add_table(rows=5, cols=2)
    t12.style = "Table Grid"
    header_row(t12, "Dimension", "Result")
    summary = [
        ("Market Heat Index",        "40.6 — Supply-heavy market (supply > demand right now)"),
        ("Circle Premium",           "100 — Market trades at 2.27× govt floor (strong real demand)"),
        ("Buy / Watch / Avoid",      "WATCH (65.4)  →  BUY (71.2) if negotiated to fair value"),
        ("5-Year ROI",               "STRONG — 70.3 % total, 11.2 % CAGR, 1.70× payoff multiple"),
    ]
    for i, (dim, res) in enumerate(summary, start=1):
        data_row(t12, i, dim, res, alt=(i % 2 == 0), bold_first=True)
    set_cell_borders(t12)

    add_note(doc,
        "The ROI is driven by price appreciation (₹9,500 → ₹16,695/sqft over 5 years). "
        "Even if appreciation is 20 % lower (exit at ₹13,356/sqft), ROI drops to ~32 % "
        "— still a MODERATE return.")

    # ──────────────────────────────────────────────────────────
    # Footer note
    # ──────────────────────────────────────────────────────────
    doc.add_paragraph()
    p_foot = doc.add_paragraph()
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_foot = p_foot.add_run(
        "Generated by GeoAI Real Estate Platform · All forecasts are model estimates, "
        "not guaranteed returns · May 2026")
    r_foot.font.size = Pt(7.5)
    r_foot.italic = True
    r_foot.font.color.rgb = MID_GREY

    # ── Save ───────────────────────────────────────────────────
    out_path = (
        r"C:\Users\kushp\OneDrive\Desktop\geoai_ml\real_estate"
        r"\Decision_Intelligence_Formula_Reference.docx"
    )
    doc.save(out_path)
    print(f"Saved → {out_path}")


if __name__ == "__main__":
    build()
